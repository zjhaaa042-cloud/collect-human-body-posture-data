from pathlib import Path
from datetime import datetime
import json

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "doc"

SOFTWARE_NAME = "人体体态数据采集系统"
SOFTWARE_SHORT = "体态数据采集系统"
VERSION = "V1.0"
FINISH_DATE = "2026年6月24日"
GENERATED_DATE = datetime.now().strftime("%Y年%m月%d日")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, header_rows=0):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(10.5)
            if i < header_rows:
                set_cell_shading(cell, "D9EAF7")


def setup_doc(title=None, landscape=False):
    doc = Document()
    sec = doc.sections[0]
    if landscape:
        sec.orientation = 1
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        st = styles[name]
        st.font.name = "黑体" if name != "Title" else "宋体"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if name != "Title" else "宋体")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(22)
        doc.add_paragraph()
    return doc


def add_para(doc, text="", first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_chapter(doc, text):
    doc.add_page_break()
    return add_heading(doc, text, 1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.35
        r = p.add_run(item)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(10.5)


def add_figure(doc, img_name, caption, width_cm=15):
    img = ROOT / img_name
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Cm(width_cm))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(9)


def add_footer_page_number(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("第 ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(t)
        fld.append(r)
        p._p.append(fld)
        p.add_run(" 页")


def make_info_doc():
    doc = setup_doc("软件著作权登记申请信息表")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("（依据项目文件整理，缺失身份类信息请在提交前补充）").font.size = Pt(10.5)

    rows = [
        ("软件全称", SOFTWARE_NAME),
        ("软件简称", SOFTWARE_SHORT),
        ("版本号", VERSION),
        ("软件分类", "应用软件 / 数据采集与处理软件"),
        ("开发完成日期", FINISH_DATE),
        ("首次发表状态", "未发表（如已公开发布，请改填首次发表日期、地点及方式）"),
        ("开发方式", "独立开发"),
        ("权利取得方式", "原始取得"),
        ("权利范围", "全部权利"),
        ("著作权人", "待填写（单位或个人全称）"),
        ("证件类型及号码", "待填写"),
        ("联系人及电话", "待填写"),
        ("电子邮箱", "待填写"),
        ("通讯地址", "待填写"),
        ("软件运行环境", "Windows 10 及以上；Python 3.10+；Node.js 18+；奥比中光 Gemini 336L 深度相机；Vosk 中文语音模型。"),
        ("主要开发语言", "Python、JavaScript、HTML、CSS"),
        ("主要技术框架", "React、Electron、Ant Design、WebSocket、OpenCV、NumPy、pyorbbecsdk、Vosk、edge-tts、Electron preload 安全桥、令牌鉴权机制。"),
        ("硬件环境", "x86/x64 计算机、USB 3.0 接口、奥比中光 Gemini 336L 深度相机或兼容深度相机。"),
        ("软件用途", "用于人体体态数据采集、RGB 图像预览、深度数据记录、点云文件生成、采集会话管理及语音辅助控制。"),
        ("技术特点", "系统采用前后端分离架构，后端负责相机流控制、深度分析、点云生成、文件保存和通信鉴权，前端通过 WebSocket 实时显示 RGB/深度预览和距离状态，并通过 Electron preload 暴露受控桌面能力。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_shading(table.cell(i, 0), "D9EAF7")
        set_cell_text(table.cell(i, 1), v)
    style_table(table)

    add_heading(doc, "软件功能概述", 1)
    add_para(doc, "人体体态数据采集系统面向姿态研究、健康监测、动作采集和三维数据建模等场景，提供采集前距离判断、实时画面预览、单次采集、会话延续、历史记录预览、语音控制、数据质量检查和多格式数据保存等能力。")
    add_heading(doc, "申请材料说明", 1)
    add_bullets(doc, [
        "说明书材料：已生成《软件操作或设计说明书》，内容不少于 10 页。",
        "源代码材料：已生成《源代码》，按项目源码行数不足 3000 行的情形提供全部源代码。",
        "申请表中的著作权人、证件号、联系方式、发表状态等身份或事实信息，需要申请人最终确认后提交。",
    ])
    doc.save(OUT / "软件登记申请信息表-城院-完成版.docx")


def make_manual_doc():
    doc = setup_doc(f"{SOFTWARE_NAME}{VERSION} 软件操作或设计说明书")
    add_para(doc, f"文档版本：{VERSION}")
    add_para(doc, f"生成日期：{GENERATED_DATE}")
    add_para(doc, "本文档依据项目目录 body_posture_collector 中的 README、INSTALL、配置文件、前端界面文件和后端核心模块整理，适用于软件著作权登记材料提交。")
    doc.add_page_break()

    add_heading(doc, "一、软件概述", 1)
    add_para(doc, "人体体态数据采集系统是一套基于深度相机的人体体态多模态数据采集工具。软件通过奥比中光 Gemini 336L 深度相机获取 RGB 图像和深度图像，并结合相机内参生成三维点云文件。系统面向人体姿态研究、康复训练数据采集、动作样本整理和三维数据建模等使用场景，能够帮助操作人员以统一流程完成数据采集与管理。")
    add_para(doc, "软件采用 Python 后端服务和 React/Electron 前端工作台组合。后端负责深度相机初始化、流式帧获取、人体距离分析、点云生成、数据质量检查、文件落盘、令牌鉴权和 WebSocket 通信；前端负责实时预览、采集控制、会话管理、历史记录查看、状态提示和用户交互，Electron 桌面端通过 preload 安全桥暴露版本、路径、鉴权令牌和关闭窗口等受控能力。")
    add_heading(doc, "1.1 主要功能", 2)
    add_bullets(doc, [
        "实时预览 RGB 彩色画面和深度画面，辅助采集人员确认被采集对象位置。",
        "基于深度图像识别人体区域并计算采集距离，给出过近、适中、过远等状态提示。",
        "支持按会话组织采集任务，新建、选择、完成采集会话，并自动维护采集元数据。",
        "支持保存 RGB 图像、深度 NPZ 数据、深度可视化 PNG 和 PLY 点云文件。",
        "支持语音识别指令和语音播报，降低采集过程中的鼠标键盘操作频率。",
        "提供历史采集列表和采集图片预览，便于对采集结果进行初步核验。",
        "支持基于一次性令牌的 WebSocket 连接鉴权，并对会话名、文件名和文本消息进行长度及字符校验。",
    ])
    add_heading(doc, "1.2 软件特点", 2)
    add_para(doc, "系统针对体态采集工作中的距离判断、文件组织和多人次数据管理需求进行设计。采集过程以会话为单位，单次采集产生的多模态文件按照固定目录结构保存，便于后续算法训练、人工标注或数据归档。")

    add_chapter(doc, "二、运行环境")
    tbl = doc.add_table(rows=1, cols=2)
    set_cell_text(tbl.cell(0, 0), "项目", True)
    set_cell_text(tbl.cell(0, 1), "要求")
    for k, v in [
        ("操作系统", "Windows 10 及以上版本"),
        ("Python 环境", "Python 3.10+，安装 requirements.txt 中列出的依赖"),
        ("前端环境", "Node.js 18+，npm 安装前端依赖"),
        ("深度相机", "奥比中光 Gemini 336L 深度相机，需正确连接并安装驱动/SDK"),
        ("语音模型", "Vosk 中文语音模型，推荐放置于 models/vosk-model-cn 或配置文件指定目录"),
        ("网络端口", "本机 WebSocket 端口 8765，前端开发服务端口 3000"),
    ]:
        row = tbl.add_row().cells
        set_cell_text(row[0], k, True)
        set_cell_text(row[1], v)
    style_table(tbl, 1)

    add_chapter(doc, "三、安装部署")
    add_heading(doc, "3.1 安装 Python 依赖", 2)
    add_numbered(doc, [
        "进入软件根目录 body_posture_collector。",
        "执行 python -m venv venv 创建虚拟环境。",
        "在 Windows 中执行 .\\venv\\Scripts\\activate 激活虚拟环境。",
        "执行 pip install -r requirements.txt 安装后端依赖。",
        "如需连接真实相机，执行 pip install pyorbbecsdk2 或从 pyorbbecsdk-v2-main 源码安装。",
    ])
    add_heading(doc, "3.2 安装前端依赖", 2)
    add_numbered(doc, [
        "进入 frontend 目录。",
        "执行 npm install 安装 React、Electron、Ant Design 等依赖。",
        "安装完成后返回项目根目录，使用启动脚本或分别启动前后端服务。",
    ])
    add_heading(doc, "3.3 配置语音模型", 2)
    add_para(doc, "软件默认启用语音模块。使用前需下载 Vosk 中文模型并解压到 models 目录，或在 config.json/config.example.json 中修改 voice.model_path。若暂不使用语音控制，可将 voice.enabled 设置为 false。")

    add_chapter(doc, "四、启动与退出")
    add_heading(doc, "4.1 一键启动", 2)
    add_para(doc, "双击或在命令行执行 go.bat。该脚本会启动后端 WebSocket 服务和前端界面。后端默认地址为 ws://localhost:8765，前端默认访问地址为 http://localhost:3000。后端启动时生成本次运行的通信令牌，并写入项目根目录 .ws_token 文件，Electron 前端通过 preload 安全桥读取该令牌后完成 WebSocket 鉴权。")
    add_heading(doc, "4.2 手动启动", 2)
    add_numbered(doc, [
        "执行 run.bat 启动后端服务。",
        "执行 run_frontend.bat 启动前端界面。",
        "也可进入 frontend 目录执行 npm run electron:dev，以 Electron 桌面窗口方式运行。",
    ])
    add_heading(doc, "4.3 退出软件", 2)
    add_para(doc, "在前端界面点击右上角退出按钮，确认后系统会向后端发送 exit_app 指令，后端停止预览、释放相机资源并关闭相关进程。")

    add_chapter(doc, "五、界面说明")
    add_figure(doc, "frontend-fixed.png", "图 1 系统主界面", 15)
    add_para(doc, "系统主界面由顶部标题栏、左侧预览区、右侧控制区和底部状态栏组成。左侧主要显示 RGB 画面、深度画面和距离提示，右侧用于会话创建、采集参数选择、单次采集、历史记录查看和语音状态显示。")
    add_heading(doc, "5.1 顶部标题栏", 2)
    add_para(doc, "顶部区域显示软件名称、辅助标题、设置按钮、帮助按钮和退出按钮。设置按钮预留系统参数配置入口，退出按钮用于关闭当前采集服务。")
    add_heading(doc, "5.2 实时预览区", 2)
    add_para(doc, "预览区通过 WebSocket 接收后端压缩后的 JPEG 预览帧。RGB 预览用于判断人物姿态与画面构图，深度预览用于判断深度数据是否有效。系统还会依据深度 ROI 或人体区域计算距离状态。")
    add_figure(doc, "distance-check.png", "图 2 距离状态检测示意", 15)
    add_heading(doc, "5.3 控制与历史区", 2)
    add_para(doc, "控制区提供采集按钮、会话选择、采集数量、最近记录和图片预览功能。用户可以新建会话后开始采集，也可以选择已有会话继续追加数据。")
    add_figure(doc, "distance-area.png", "图 3 距离检测区域示意", 13)

    add_chapter(doc, "六、操作流程")
    add_heading(doc, "6.1 新建采集会话", 2)
    add_numbered(doc, [
        "启动后端服务并打开前端界面。",
        "确认底部状态栏显示已连接。",
        "在会话区域输入会话名称，名称可包含中文、英文、数字、下划线和短横线。",
        "点击新建会话，系统在 data/sessions 目录下建立对应文件夹，并初始化 metadata.json。",
    ])
    add_heading(doc, "6.2 调整采集距离", 2)
    add_numbered(doc, [
        "被采集人员站在相机前方。",
        "观察界面距离状态，系统默认目标距离为 1000mm，容差为 200mm。",
        "当状态提示过近时后退，当提示过远时前移。",
        "距离状态为合适后保持姿态，准备采集。",
    ])
    add_heading(doc, "6.3 执行单次采集", 2)
    add_numbered(doc, [
        "在控制区选择是否保存 RGB、深度数据和点云数据。",
        "点击采集按钮或发出“开始采集”语音指令。",
        "系统进行短暂等待，获取当前帧并检测人体区域。",
        "若采集成功，系统保存文件并更新历史记录；若失败，界面显示失败原因。",
    ])
    add_heading(doc, "6.4 查看采集历史", 2)
    add_para(doc, "采集完成后，右侧历史区域会显示最近采集记录。对于包含 RGB 图像的记录，用户可以点击预览，系统读取对应文件并以图片弹窗展示。")
    add_heading(doc, "6.5 完成会话", 2)
    add_para(doc, "采集任务结束后点击完成会话或使用“完成”语音指令。系统将当前会话状态写入 metadata.json，并记录完成时间。")

    add_chapter(doc, "七、数据保存说明")
    add_para(doc, "系统以会话为单位保存采集结果。默认输出目录为 data/sessions，每个会话目录下包含 rgb、depth、pointcloud 三个子目录和 metadata.json 元数据文件。文件管理模块还预留 exports 导出目录和 config 配置目录，可用于会话压缩导出、会话删除、存储统计和配置读写等扩展管理能力。")
    t = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["目录/文件", "格式", "内容", "用途"]):
        set_cell_text(t.cell(0, i), h, True)
    for rowdata in [
        ("rgb", "PNG", "彩色图像文件", "用于姿态观察、人工核验和后续图像算法处理"),
        ("depth", "NPZ/PNG", "原始深度矩阵和深度可视化图", "用于距离计算、三维重建和深度质量分析"),
        ("pointcloud", "PLY", "三维点云数据", "用于三维建模、空间分析和可视化"),
        ("metadata.json", "JSON", "会话信息、相机参数、采集统计、文件索引", "用于追踪每次采集结果和会话状态"),
    ]:
        cells = t.add_row().cells
        for i, v in enumerate(rowdata):
            set_cell_text(cells[i], v, i == 0)
    style_table(t, 1)

    add_chapter(doc, "八、核心模块设计")
    modules = [
        ("相机管理模块", "backend/core/camera_manager.py", "负责深度相机初始化、流启动/停止、彩色帧和深度帧读取、相机内参获取、点云数据生成以及无真实设备时的模拟数据生成。"),
        ("深度分析模块", "backend/core/depth_analyzer.py", "对深度图进行人体区域筛选、形态学处理、轮廓定位和距离状态判断，输出距离数值、状态、置信度和提示文本。"),
        ("数据采集模块", "backend/core/data_collector.py", "负责会话目录创建、采集质量检查、RGB/深度/点云保存、采集编号生成和元数据维护。"),
        ("WebSocket 服务模块", "backend/server/ws_server.py", "负责前后端实时通信，处理预览、采集、会话、历史、语音和退出等消息；启动时生成通信令牌，提供 /auth-token 查询接口，并要求客户端连接后先发送 auth 鉴权消息。"),
        ("文件管理模块", "backend/storage/file_manager.py", "负责会话目录、导出目录和配置目录管理，提供会话元数据读写、会话压缩导出、会话删除、存储容量统计和配置文件读写能力。"),
        ("语音控制模块", "backend/voice", "负责语音识别、指令解析和语音播报，支持开始采集、停止、下一个和完成等指令。"),
        ("前端界面模块", "frontend/src", "负责主界面布局、预览展示、采集控制、会话列表、历史记录、图片预览和状态栏显示。"),
        ("Electron 桌面模块", "frontend/electron-main.js、frontend/preload.js", "负责桌面窗口创建、React 服务启动、受控 IPC 接口暴露、通信令牌读取和窗口关闭等桌面端能力；采用 contextIsolation、sandbox 和 preload 降低渲染进程直接访问系统资源的风险。"),
    ]
    mt = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["模块", "对应文件", "功能说明"]):
        set_cell_text(mt.cell(0, i), h, True)
    for rowdata in modules:
        cells = mt.add_row().cells
        for i, v in enumerate(rowdata):
            set_cell_text(cells[i], v, i == 0)
    style_table(mt, 1)

    add_chapter(doc, "九、参数配置")
    add_para(doc, "软件可通过配置文件调整相机、语音、存储、距离和界面参数。config.example.json 给出了默认配置示例。")
    ct = doc.add_table(rows=1, cols=3)
    for i, h in enumerate(["配置项", "默认值", "说明"]):
        set_cell_text(ct.cell(0, i), h, True)
    for rowdata in [
        ("camera.width / height", "1280 / 800", "相机采集分辨率"),
        ("camera.fps", "30", "相机采集帧率"),
        ("storage.output_dir", "data", "数据输出目录"),
        ("storage.save_rgb", "true", "是否保存 RGB 图像"),
        ("storage.save_depth", "true", "是否保存深度数据"),
        ("storage.save_pointcloud", "true", "是否保存点云数据"),
        ("distance.target_distance_mm", "1000", "推荐采集距离"),
        ("distance.tolerance_mm", "200", "距离容差范围"),
        ("gui.preview_fps", "20", "前端预览刷新帧率"),
        ("websocket_port", "8765", "后端 WebSocket 服务端口"),
    ]:
        cells = ct.add_row().cells
        for i, v in enumerate(rowdata):
            set_cell_text(cells[i], v, i == 0)
    style_table(ct, 1)

    add_chapter(doc, "十、异常处理与维护")
    add_bullets(doc, [
        "相机未检测到：检查 USB 连接、设备管理器、相机驱动和 pyorbbecsdk 安装状态。",
        "前端无法连接：确认后端服务已启动，检查 8765 端口是否被占用或被安全软件拦截。",
        "连接鉴权失败：确认后端服务已重新启动并生成 .ws_token，Electron 前端可通过 preload 读取令牌；浏览器调试模式下可通过 /auth-token 获取本次运行令牌。",
        "语音模型加载失败：确认模型文件完整，并检查 voice.model_path 配置是否正确。",
        "采集失败：检查画面亮度、深度覆盖率、被采集对象是否位于相机前方，以及磁盘剩余空间是否充足。",
        "图片无法预览：确认当前会话下 rgb 目录存在对应 PNG 文件，且文件名未被手工修改。",
    ])

    add_chapter(doc, "十一、安全与数据管理")
    add_para(doc, "软件默认在本机目录保存采集数据，不主动上传外部服务器。采集数据可能包含人体图像和体态信息，使用单位应按照数据安全和个人信息保护要求，明确采集授权、使用范围、保存期限和访问权限。")
    add_para(doc, "当前版本在通信层加入运行期令牌鉴权。后端启动时生成随机令牌，写入 .ws_token 文件并通过本机 /auth-token 接口供前端获取；WebSocket 客户端连接后需在限定时间内发送 auth 消息，令牌不匹配则关闭连接。后端还对会话名称、文件名和播报文本进行长度和字符校验，减少非法路径、异常参数和非预期输入带来的风险。")
    add_para(doc, "Electron 桌面端关闭 Node 集成，启用 contextIsolation 和 sandbox，通过 preload.js 暴露有限的 electronAPI 接口，避免渲染界面直接访问 Node.js 系统能力。退出软件时，后端会停止预览、释放相机资源，并按端口清理前端开发服务进程。")

    add_chapter(doc, "十二、附录：软件目录结构")
    add_para(doc, "backend：后端核心代码目录；frontend：前端界面、Electron 入口及 preload 安全桥目录；config：相机参数目录；data：默认采集数据目录；models：语音模型目录；logs：运行日志目录；.ws_token：后端运行期生成的 WebSocket 通信令牌文件；requirements.txt：后端依赖列表；go.bat、run.bat、run_frontend.bat：启动脚本。")
    add_footer_page_number(doc)
    doc.save(OUT / "软件操作或设计说明书-10页以上-完成版.docx")


def collect_source_files():
    include_ext = {".py", ".js", ".css", ".html", ".json", ".bat"}
    exclude_parts = {
        ".git", "node_modules", "__pycache__", ".mimocode", ".playwright-mcp",
        ".serena", "output", "tmp", "temp", "Log", "logs", "data", "models", "tools"
    }
    exclude_names = {"package-lock.json"}
    files = []
    for p in sorted(ROOT.rglob("*")):
        rel = p.relative_to(ROOT)
        if not p.is_file():
            continue
        if p.suffix.lower() not in include_ext:
            continue
        if any(part in exclude_parts for part in rel.parts):
            continue
        if p.name in exclude_names:
            continue
        text = p.read_text(encoding="utf-8", errors="replace")
        files.append((rel.as_posix(), text))
    return files


def make_source_doc():
    files = collect_source_files()
    total_lines = sum(len(text.splitlines()) for _, text in files)
    source_note = (
        "因项目源代码不足 3000 行，本文件提供全部源代码。"
        if total_lines < 3000
        else "因项目源代码超过 3000 行，本文件提供全部项目实现源码，满足不少于 3000 行的提交要求。"
    )
    doc = setup_doc(f"{SOFTWARE_NAME}{VERSION} 源代码", landscape=True)
    add_para(doc, f"软件名称：{SOFTWARE_NAME}")
    add_para(doc, f"版本号：{VERSION}")
    add_para(doc, f"源码统计：共 {len(files)} 个源文件，{total_lines} 行。{source_note}")
    add_para(doc, "说明：本源码文档不包含 node_modules、运行日志、采集数据、语音模型、临时文件、构建缓存、说明文档、生成工具及前端 package-lock.json 等自动生成、第三方依赖或非实现文件。")
    doc.add_page_break()

    code_style = doc.styles.add_style("CodeBlock", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    code_style.font.size = Pt(7.5)
    code_style.paragraph_format.line_spacing = 1.0
    code_style.paragraph_format.space_after = Pt(0)

    line_no = 1
    for rel, text in files:
        add_heading(doc, f"文件：{rel}", 1)
        for raw in text.splitlines():
            p = doc.add_paragraph(style="CodeBlock")
            run = p.add_run(f"{line_no:04d}  {raw}")
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(7.5)
            line_no += 1
        doc.add_paragraph()
    add_footer_page_number(doc)
    doc.save(OUT / "源代码-实现具体功能有开头结尾3000行以上(不到3000行则提供全部源代码)-完成版.docx")
    return total_lines, len(files)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    make_info_doc()
    make_manual_doc()
    total_lines, file_count = make_source_doc()
    summary = {
        "generated_at": datetime.now().isoformat(),
        "output_dir": str(OUT),
        "software_name": SOFTWARE_NAME,
        "version": VERSION,
        "source_file_count": file_count,
        "source_total_lines": total_lines,
    }
    (OUT / "生成说明.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
